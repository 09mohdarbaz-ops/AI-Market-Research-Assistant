from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from textwrap import wrap

import streamlit as st


SECTIONS = [
    "Executive Summary",
    "SWOT",
    "PESTLE",
    "Porter's Five Forces",
    "Market Segmentation",
    "Future Outlook",
]


@dataclass(frozen=True)
class ReportInputs:
    industry: str
    country: str
    forecast_period: str


def clean_input(value: str, fallback: str) -> str:
    value = value.strip()
    return value if value else fallback


def build_report(inputs: ReportInputs) -> dict[str, list[str]]:
    industry = inputs.industry
    country = inputs.country
    period = inputs.forecast_period

    return {
        "Executive Summary": [
            f"The {industry} industry in {country} is expected to evolve across the {period} forecast period as demand patterns, investment priorities, regulation, and competitive positioning continue to shift.",
            "Key themes to monitor include margin resilience, supply chain stability, technology adoption, customer affordability, and the ability of leading players to convert market intelligence into faster execution.",
            "This report provides a structured strategic view intended for planning discussions, early market sizing work, and executive decision support.",
        ],
        "SWOT": [
            f"Strengths: Established demand pools, local operating knowledge, and opportunities to tailor {industry.lower()} offerings to {country}'s buyer needs.",
            "Weaknesses: Cost volatility, capability gaps, fragmented distribution, and uneven access to reliable market data may slow execution.",
            "Opportunities: Digital channels, partnerships, premium segments, underserved regions, and productivity improvements can unlock new growth.",
            "Threats: Regulatory change, aggressive entrants, substitution, currency pressure, and slower macroeconomic activity could pressure growth assumptions.",
        ],
        "PESTLE": [
            f"Political: Public policy priorities and trade relationships will influence the operating climate for {industry.lower()} companies in {country}.",
            "Economic: Inflation, interest rates, household or enterprise spending power, and investment cycles will shape demand.",
            "Social: Demographic shifts, trust, convenience expectations, and sustainability preferences may redefine product and service design.",
            "Technological: Automation, analytics, AI-enabled workflows, and platform integration can improve speed, personalization, and cost discipline.",
            "Legal: Compliance obligations, licensing, consumer protection, labor rules, and data governance require active monitoring.",
            "Environmental: Energy use, resource efficiency, emissions reporting, and circularity pressures are likely to become more material.",
        ],
        "Porter's Five Forces": [
            "Competitive rivalry: Rivalry is likely to intensify where differentiation is limited and switching costs are low.",
            "Threat of new entrants: Entry barriers depend on capital intensity, licensing, distribution access, brand trust, and specialist talent.",
            "Buyer power: Customers gain leverage when alternatives are easy to compare or when procurement is centralized.",
            "Supplier power: Supplier concentration, imported inputs, scarce skills, and logistics constraints can affect pricing power.",
            "Threat of substitutes: Adjacent categories, technology-enabled alternatives, and changing user habits can redirect demand.",
        ],
        "Market Segmentation": [
            "Customer type: Segment by enterprise, small business, public sector, and consumer demand where relevant.",
            "Product or service tier: Separate value, mainstream, premium, and specialist propositions to clarify margin and volume tradeoffs.",
            "Channel: Compare direct sales, distributor-led models, marketplaces, partnerships, and digital self-service channels.",
            "Geography: Distinguish mature urban demand from regional or underserved growth pockets within the country.",
            "Use case: Group demand by mission-critical, discretionary, compliance-driven, and replacement-led buying behavior.",
        ],
        "Future Outlook": [
            f"Across {period}, the {industry} market in {country} is likely to reward companies that combine disciplined pricing with targeted innovation and resilient operations.",
            "Near-term planning should focus on demand validation, competitor benchmarking, regulatory watchlists, and channel productivity.",
            "Longer-term advantage will come from data-rich customer relationships, scalable delivery models, and the ability to respond quickly as market signals change.",
        ],
    }


def render_report(report: dict[str, list[str]]) -> None:
    for section in SECTIONS:
        st.header(section)
        for paragraph in report[section]:
            st.write(paragraph)


def create_word_document(inputs: ReportInputs, report: dict[str, list[str]]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("Install python-docx to enable Word export.") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_heading(f"{inputs.industry} Market Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"{inputs.country} | Forecast period: {inputs.forecast_period}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section_name in SECTIONS:
        document.add_heading(section_name, level=1)
        for paragraph in report[section_name]:
            document.add_paragraph(paragraph, style="Body Text")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def create_pdf_document(inputs: ReportInputs, report: dict[str, list[str]]) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("Install reportlab to enable PDF export.") from exc

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        rightMargin=0.72 * inch,
        leftMargin=0.72 * inch,
        topMargin=0.68 * inch,
        bottomMargin=0.68 * inch,
        title=f"{inputs.industry} Market Report",
    )
    styles = getSampleStyleSheet()
    styles["Title"].textColor = colors.HexColor("#1F2937")
    styles["Heading1"].textColor = colors.HexColor("#0F766E")
    styles["BodyText"].leading = 14

    story = [
        Paragraph(f"{inputs.industry} Market Report", styles["Title"]),
        Paragraph(f"{inputs.country} | Forecast period: {inputs.forecast_period}", styles["BodyText"]),
        Spacer(1, 0.2 * inch),
    ]

    for section_name in SECTIONS:
        story.append(Paragraph(section_name, styles["Heading1"]))
        for paragraph in report[section_name]:
            story.append(Paragraph(paragraph, styles["BodyText"]))
            story.append(Spacer(1, 0.07 * inch))
        story.append(Spacer(1, 0.12 * inch))

    doc.build(story)
    return output.getvalue()


def filename_for(inputs: ReportInputs, extension: str) -> str:
    stem = f"{inputs.industry}_{inputs.country}_{inputs.forecast_period}".lower()
    safe = "".join(char if char.isalnum() else "_" for char in stem)
    while "__" in safe:
        safe = safe.replace("__", "_")
    return f"{safe.strip('_')}_market_report.{extension}"


def main() -> None:
    st.set_page_config(page_title="Market Intelligence Report Builder", layout="wide")

    st.markdown(
        """
        <style>
            .block-container { padding-top: 2rem; max-width: 1120px; }
            h1, h2, h3 { letter-spacing: 0; }
            section[data-testid="stSidebar"] { background: #f8fafc; }
            .stDownloadButton button { width: 100%; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.title("Report Inputs")
        industry = st.text_input("Industry", value="Renewable Energy")
        country = st.text_input("Country", value="United States")
        forecast_period = st.text_input("Forecast Period", value="2026-2031")
        generate = st.button("Generate report", type="primary", use_container_width=True)

    inputs = ReportInputs(
        industry=clean_input(industry, "Selected Industry"),
        country=clean_input(country, "Selected Country"),
        forecast_period=clean_input(forecast_period, "Selected Period"),
    )
    report = build_report(inputs)

    st.title("Market Intelligence Report Builder")
    st.caption("Generate a structured strategy report and export it as Word or PDF.")

    if generate:
        st.success("Report generated.")

    top_cols = st.columns([1, 1, 2])
    with top_cols[0]:
        try:
            st.download_button(
                "Export Word",
                data=create_word_document(inputs, report),
                file_name=filename_for(inputs, "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except RuntimeError as exc:
            st.warning(str(exc))
    with top_cols[1]:
        try:
            st.download_button(
                "Export PDF",
                data=create_pdf_document(inputs, report),
                file_name=filename_for(inputs, "pdf"),
                mime="application/pdf",
                use_container_width=True,
            )
        except RuntimeError as exc:
            st.warning(str(exc))

    st.divider()
    render_report(report)


if __name__ == "__main__":
    main()
